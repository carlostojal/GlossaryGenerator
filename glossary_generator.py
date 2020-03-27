
#
# Copyright (c) Carlos Tojal 2020
# GlossaryGenerator
# glossary_generator.py
#

from docx import Document
import requests
import json
from io import BytesIO

print("\n** GlossaryGenerator **")
print("Developed by Carlos Tojal\n")

# loads configuration from file "config.json"
print("Loading configuration...")
f = open("config.json", "r")
config = json.loads(f.read())
f.close()
print("Loaded configuration.\n")

# loads terms from file "terms.txt"
print("Loading terms...")
f = open("terms.txt", "r")
terms = f.read().split("\n")
f.close()
print("Loaded terms.\n")

filename = config['output_path'] + "/" + config['filename']

# creates files requested in configuration
if config['docx'] == "true":
    document = Document()

    # adds title do document
    document.add_heading(config['title'], level = 0)

    document.add_paragraph() # blank line

if config['json'] == "true":
    glossary = {} # stores what will be outputted to the JSON file

if config['csv'] == "true":
    csv = "" # stores what will be outputted to the CSV file

if config['txt'] == "true":
    txt = "" # stores what will be outputted to the TXT file

print("Generating glossary...\n")

print("Searching for terms...")

for term in terms:
    print("\nSearching for term \"" + term + "\"...")
    response = requests.get("https://api.duckduckgo.com/?q=" + term + "&format=json&pretty=1") # request DuckDuckGo API to get term definition
    search = json.loads(response.text)
    raw_text = search['AbstractText']
    raw_text = raw_text.split(".")
    text = ""
    x = 0
    # add requested number of phrase, or the maximum available in the result
    for phrase in raw_text:
        if x < int(config['n_phrases']):
            text += phrase
            if x < len(raw_text) - 1: # to not write the point in the last phrase
                text += "."
        x += 1
    if config['docx'] == "true":
        p = document.add_paragraph(style = "List Bullet")
        r = p.add_run(term + ": ")
        r.bold = True # term is bold in document
    if config['csv'] == "true":
        csv += term + ";"
    if config['txt'] == "true":
        txt += term + ": "
    if text != "": # if text was returned
        print("Found.")
        if config['docx'] == "true":
            r = p.add_run(text)
            if config['insert_term_image'] == "true": # if the user chose to insert term image
                if search['Image'] != "": # if an image is available
                    # get image from URL and insert to document
                    response1 = requests.get(search['Image'])
                    image = BytesIO(response1.content)
                    document.add_picture(image)
        if config['json'] == "true":
            glossary[term] = text
        if config['csv'] == "true":
            csv += text + "\n"
        if config['txt'] == "true":
            txt += text + "\n"
    else:
        print("Not found.")
        if config['docx'] == "true":
            r = p.add_run("Not found.")
        if config['json'] == "true":
            glossary[term] = "Not found."
        if config['csv'] == "true":
            csv += "Not found.\n"
            f = open(filename + ".csv", "a")
            f.write("Not found.\n")
            f.close()
        if config['txt'] == "true":
            txt += "Not found.\n"
            f = open(filename + ".txt", "a")
            f.write("Not found.\n")
            f.close()

print("\nGenerated glossary.")

# writes header, footer and credits to DOCX document
if config['docx'] == "true":
    print("\nWriting header...")
    # writes header loaded from configuration
    header = document.sections[0].header
    p = header.paragraphs[0]
    p.text = config['header_text']
    print("Written header.")

    print("\nWriting footer...")
    # writes footer loaded from configuration
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.text = config['footer_text']
    print("Written footer.")
    
    if config['credits'] == "true": # append credits to footer if user chose that option
        footer = document.sections[0].footer
        p = footer.paragraphs[0]
        p.text += "\n\nGenerated using GlossaryGenerator\n(https://github.com/carlostojal/GlossaryGenerator)"

    # saves document
    print("\nSaving document \"" + filename + ".docx\"...")
    document.save(filename + ".docx")
    print("Saved document \"" + filename + ".docx\".")

# Adds credits to glossary object and saves to JSON file
if config['json'] == "true":
    if config['credits'] == "true":
        glossary['Credits'] = "Generated using GlossaryGenerator (https://github.com/carlostojal/GlossaryGenerator)"
    print("\nSaving document \"" + filename + ".json\"...")
    f = open(filename + ".json", "w")
    f.write(json.dumps(glossary, indent = 4, sort_keys = True))
    f.close()
    print("Saved document \"" + filename + ".json\".")

# Adds credits to the variable and saves to CSV file
if config['csv'] == "true":
    if config['credits'] == "true":
        csv += "Credits;Generated using GlossaryGenerator (https://github.com/carlostojal/GlossaryGenerator)"
    print("\nSaving document \"" + filename + ".csv\"...")
    f = open(filename + ".csv", "w")
    f.write(csv)
    f.close()
    print("\nSaved document \"" + filename + ".csv\".")

# Adds credits to the variable and saves to TXT file
if config['txt'] == "true":
    if config['credits'] == "true":
        txt += "Credits: Generated using GlossaryGenerator (https://github.com/carlostojal/GlossaryGenerator)"
    print("\nSaving document \"" + filename + ".txt\"...")
    f = open(filename + ".txt", "w")
    f.write(txt)
    f.close()
    print("\nSaved document \"" + filename + ".txt\".")

